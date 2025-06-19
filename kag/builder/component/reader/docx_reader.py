# -*- coding: utf-8 -*-
# Copyright 2023 OpenSPG Authors
#
# Licensed under the Apache License, Version 2.0 (the "License"); you may not use this file except
# in compliance with the License. You may obtain a copy of the License at
#
# http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software distributed under the License
# is distributed on an "AS IS" BASIS, WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express
# or implied.

import os
from typing import List, Union, Dict, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from kag.interface import LLMClient
from kag.builder.model.chunk import Chunk, ChunkTypeEnum
from kag.interface import ReaderABC
from kag.builder.prompt.outline_prompt import OutlinePrompt
from kag.common.utils import generate_hash_id
from knext.common.base.runnable import Input, Output
from kag.builder.model.sub_graph import SubGraph, Node, Edge
from kag.builder.component.splitter.length_splitter import LengthSplitter


class DocxNode:
    def __init__(
        self, title: str, level: int, content: str = "", node_type: str = "heading"
    ):
        """
        Initialize a DocxNode.

        Args:
            title (str): The title/text of the node
            level (int): The hierarchical level
            content (str): The content text
            node_type (str): Type of node ("heading", "list", "paragraph", etc.)
        """
        self.title = title
        self.display_title = (
            title[:100] + "..." if len(title) > 100 else title
        )  # Truncated for display
        self.level = level
        self.content = content
        self.node_type = node_type
        self.children: List[DocxNode] = []
        self.properties: Dict[
            str, str
        ] = {}  # Store additional properties like style, indent level, etc.

    def __str__(self):
        return f"{self.node_type}({self.level}): {self.display_title}"

    def add_child(self, child: "DocxNode") -> None:
        """Add a child node with proper level adjustment if needed."""
        if child.level <= self.level and self.node_type == "heading":
            # Adjust child level to maintain hierarchy
            level_diff = self.level + 1 - child.level
            child.level += level_diff
        self.children.append(child)

    def get_node_id(self) -> str:
        """Generate a stable node ID that's not dependent on the full title."""
        return generate_hash_id(f"{self.node_type}_{self.level}_{self.title[:100]}")


@ReaderABC.register("docx")
@ReaderABC.register("docx_reader")
class DocxReader(ReaderABC):
    """
    A class for reading Docx files into Chunk objects.

    This class inherits from ReaderABC and provides the functionality to process Docx files,
    extract their text content, and convert it into a list of Chunk objects.
    """

    def __init__(
        self,
        llm: LLMClient = None,
        cut_depth: int = 5,
        length_splitter: LengthSplitter = None,
        **kwargs,
    ):
        """
        Initializes the DocxReader with an optional LLMClient instance.

        Args:
            llm (LLMClient): An optional LLMClient instance used for generating outlines. Defaults to None.
            cut_depth (int): The heading level at which to cut the document into chunks. Defaults to 5.
            length_splitter (LengthSplitter): An optional LengthSplitter instance for splitting long chunks. Defaults to None.
        """
        super().__init__(**kwargs)
        self.llm = llm
        self.prompt = OutlinePrompt(self.kag_project_config.language)
        self.cut_depth = cut_depth
        self.length_splitter = length_splitter

    def outline_chunk(self, chunk: Union[Chunk, List[Chunk]], basename) -> List[Chunk]:
        """
        Generates outlines for the given chunk(s) and separates the content based on these outlines.

        Args:
            chunk (Union[Chunk, List[Chunk]]): A single Chunk object or a list of Chunk objects.
            basename: The base name used for generating chunk IDs and names.

        Returns:
            List[Chunk]: A list of Chunk objects separated by the generated outlines.
        """
        if isinstance(chunk, Chunk):
            chunk = [chunk]
        outlines = []
        for c in chunk:
            outline = self.llm.invoke({"input": c.content}, self.prompt)
            outlines.extend(outline)
        content = "\n".join([c.content for c in chunk])
        chunks = self.sep_by_outline(content, outlines, basename)
        return chunks

    def sep_by_outline(self, content, outlines, basename):
        """
        Separates the content based on the provided outlines.

        Args:
            content (str): The content to be separated.
            outlines (List[str]): A list of outlines used to separate the content.
            basename: The base name used for generating chunk IDs and names.

        Returns:
            List[Chunk]: A list of Chunk objects separated by the provided outlines.
        """
        position_check = []
        for outline in outlines:
            start = content.find(outline)
            position_check.append((outline, start))
        chunks = []
        for idx, pc in enumerate(position_check):
            chunk = Chunk(
                id=generate_hash_id(f"{basename}#{pc[0]}"),
                name=f"{basename}#{pc[0]}",
                content=content[
                    pc[1] : (
                        position_check[idx + 1][1]
                        if idx + 1 < len(position_check)
                        else len(position_check)
                    )
                ],
            )
            chunks.append(chunk)
        return chunks

    @staticmethod
    def _get_heading_level(paragraph: Paragraph) -> Tuple[int, bool]:
        """
        Get the heading level of a paragraph and determine if it's a real heading.

        Args:
            paragraph (Paragraph): The paragraph to check.

        Returns:
            Tuple[int, bool]: A tuple containing:
                - The heading level (1-6) if the paragraph is a heading, 0 otherwise
                - Boolean indicating if this is a real heading (based on styles) or just formatted text
        """
        # 1. Check for built-in heading styles - Most reliable method
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            # Check for standard heading styles (Heading 1, Heading 2, etc.)
            if style_name.startswith("heading"):
                try:
                    return int(style_name[7:]), True
                except ValueError:
                    pass
            # Check for title style
            elif style_name == "title":
                return 1, True
            # Check for TOC styles (Table of Contents)
            elif style_name.startswith("toc"):
                try:
                    return int(style_name[3:]), True
                except ValueError:
                    pass
            # Check for outline level property if available
            elif (
                hasattr(paragraph.style, "outline_level")
                and paragraph.style.outline_level
            ):
                return paragraph.style.outline_level, True

        # 2. Check for outline level in paragraph properties
        if hasattr(paragraph._element, "pPr") and paragraph._element.pPr is not None:
            pPr = paragraph._element.pPr
            if (
                pPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl"
                )
                is not None
            ):
                outline_level = pPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl"
                ).get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                )
                if outline_level is not None:
                    try:
                        return (
                            int(outline_level) + 1,
                            True,
                        )  # Convert 0-based to 1-based level
                    except ValueError:
                        pass

        # 3. Check for special formatting as fallback - Less reliable
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold or (
                first_run.font.size and first_run.font.size > 12 * 20
            ):  # Size in twips (20 twips = 1 point)
                # If text is numbered (e.g., "1.", "1.1", "1.1.1")
                text = paragraph.text.strip()
                if text and text[0].isdigit():
                    # Count the number of dot-separated numbers
                    parts = text.split(" ", 1)[0].rstrip(".")
                    if all(p.isdigit() for p in parts.split(".")):
                        return len(parts.split(".")), False
                # Otherwise treat as level 1 heading
                return 1, False

        return 0, False

    def _extract_run_text(self, run) -> str:
        """
        Extract text from a run, including any special characters, symbols, and fields.

        Args:
            run: A run object from python-docx

        Returns:
            str: The extracted text
        """
        text = run.text

        # Handle special elements
        if hasattr(run, "_element") and run._element is not None:
            for elem in run._element:
                if elem.tag.endswith("tab"):
                    text += "\t"
                elif elem.tag.endswith("br"):  # Line break
                    text += "\n"
                elif elem.tag.endswith("cr"):  # Carriage return
                    text += "\n"
                elif elem.tag.endswith("noBreakHyphen"):  # Non-breaking hyphen
                    text += "-"
                elif elem.tag.endswith("softHyphen"):  # Soft hyphen
                    text += "\u00ad"
                elif elem.tag.endswith("sym"):  # Symbol
                    if "char" in elem.attrib:
                        text += elem.attrib["char"]
                elif elem.tag.endswith("fldChar"):  # Field character (complex field)
                    if elem.get("fldCharType") == "begin":
                        # Start of a complex field (like a hyperlink)
                        pass
                elif elem.tag.endswith("instrText"):  # Field instruction text
                    field_text = elem.text
                    if field_text and field_text.strip().startswith("HYPERLINK"):
                        # Extract hyperlink URL
                        url = field_text.split('"')[1] if '"' in field_text else ""
                        text += f"[{text}]({url})"
                elif elem.tag.endswith("drawing"):  # Inline image
                    text += "[IMAGE]"

        return text

    def _extract_paragraph_text(self, paragraph: Paragraph) -> str:
        """
        Extract text from a paragraph, preserving formatting and special characters.

        Args:
            paragraph: A paragraph object from python-docx

        Returns:
            str: The extracted text with formatting indicators
        """
        text_parts = []

        # Handle paragraph formatting
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if not (style_name.startswith("heading") or style_name == "title"):
                text_parts.append(f"[{style_name}]")

        # Handle paragraph alignment
        if paragraph.alignment:
            alignment_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
            if paragraph.alignment in alignment_map:
                text_parts.append(f"[align={alignment_map[paragraph.alignment]}]")

        for run in paragraph.runs:
            run_text = self._extract_run_text(run)
            if run_text.strip():
                formatted_text = run_text

                # Basic formatting
                if run.bold:
                    formatted_text = f"**{formatted_text}**"
                if run.italic:
                    formatted_text = f"*{formatted_text}*"
                if run.underline:
                    formatted_text = f"_{formatted_text}_"

                # Additional formatting - check attributes carefully
                if hasattr(run, "_element") and run._element is not None:
                    # Check for strikethrough
                    if run._element.rPr is not None:
                        strike = run._element.rPr.find(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike"
                        )
                        dstrike = run._element.rPr.find(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}dstrike"
                        )
                        if strike is not None or dstrike is not None:
                            formatted_text = f"~~{formatted_text}~~"

                # Font properties
                if hasattr(run, "font"):
                    # Vertical alignment (subscript/superscript)
                    if hasattr(run.font, "subscript") and run.font.subscript:
                        formatted_text = f"<sub>{formatted_text}</sub>"
                    if hasattr(run.font, "superscript") and run.font.superscript:
                        formatted_text = f"<sup>{formatted_text}</sup>"

                    # Text color
                    if (
                        hasattr(run.font, "color")
                        and run.font.color
                        and run.font.color.rgb
                    ):
                        formatted_text = f'<span style="color: #{run.font.color.rgb}">{formatted_text}</span>'

                    # Highlight color
                    if (
                        hasattr(run.font, "highlight_color")
                        and run.font.highlight_color
                    ):
                        formatted_text = f"<mark>{formatted_text}</mark>"

                text_parts.append(formatted_text)

        return " ".join(text_parts)

    def _extract_footnote_text(self, footnote) -> str:
        """Extract text from a footnote."""
        footnote_text = []
        for paragraph in footnote.paragraphs:
            para_text = self._extract_paragraph_text(paragraph)
            if para_text.strip():
                footnote_text.append(para_text)
        return " ".join(footnote_text)

    def _extract_comment_text(self, comment) -> str:
        """Extract text from a comment."""
        return f"[Comment by {comment.author} at {comment.date}: {comment.text}]"

    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table, preserving structure.

        Args:
            table: A table object from python-docx

        Returns:
            str: The extracted table text in a structured format
        """
        table_text = []
        for i, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                # Extract all paragraphs from the cell
                cell_text = []
                for paragraph in cell.paragraphs:
                    para_text = self._extract_paragraph_text(paragraph)
                    if para_text.strip():
                        cell_text.append(para_text)
                row_text.append(" ".join(cell_text))

            # Format row as markdown table row
            if i == 0:
                table_text.append("| " + " | ".join(row_text) + " |")
                table_text.append("| " + " | ".join(["---"] * len(row_text)) + " |")
            else:
                table_text.append("| " + " | ".join(row_text) + " |")

        return "\n".join(table_text)

    def _extract_list_text(self, paragraph: Paragraph, list_style: str = None) -> str:
        """
        Extract text from a list item, preserving list formatting.

        Args:
            paragraph: A paragraph object from python-docx
            list_style: The style of the list (bullet or number)

        Returns:
            str: The extracted list item text
        """
        text = self._extract_paragraph_text(paragraph)

        # Determine list style and level
        level = 0
        if paragraph._element.pPr and paragraph._element.pPr.numPr:
            numPr = paragraph._element.pPr.numPr
            if hasattr(numPr, "numId") and numPr.numId.val:
                # It's a numbered list
                if hasattr(numPr, "ilvl"):
                    level = int(numPr.ilvl.val)
                return f"{'    ' * level}1. {text}"

        # Default to bullet list
        return f"{'    ' * level}* {text}"

    def _build_document_tree(self, doc: Document) -> DocxNode:
        """
        Build a tree structure from the document with improved content extraction.

        Args:
            doc (Document): The document to process.

        Returns:
            DocxNode: The root node of the document tree.
        """
        root = DocxNode("root", 0, node_type="root")
        stack = [root]
        current_content = []
        in_list = False

        print(f"Processing document with {len(doc.paragraphs)} paragraphs")

        # First, collect all sections to handle headers, footers, and footnotes
        sections = doc.sections
        for section in sections:
            # Process headers
            for header in [section.header]:
                if header.is_linked_to_previous:
                    continue
                header_text = []
                for paragraph in header.paragraphs:
                    para_text = self._extract_paragraph_text(paragraph)
                    if para_text.strip():
                        header_text.append(para_text)
                if header_text:
                    current_content.append("[Header]")
                    current_content.extend(header_text)
                    current_content.append("")

            # Process footers
            for footer in [section.footer]:
                if footer.is_linked_to_previous:
                    continue
                footer_text = []
                for paragraph in footer.paragraphs:
                    para_text = self._extract_paragraph_text(paragraph)
                    if para_text.strip():
                        footer_text.append(para_text)
                if footer_text:
                    current_content.append("[Footer]")
                    current_content.extend(footer_text)
                    current_content.append("")

        # Process main document body
        for element in doc._body._body:
            try:
                if element.tag.endswith("p"):  # Paragraph
                    para = Paragraph(element, doc)
                    # Get paragraph properties
                    style_name = para.style.name if para.style else "Normal"

                    # Determine paragraph type and level
                    heading_level, is_real_heading = self._get_heading_level(para)
                    is_list = bool(para._element.pPr and para._element.pPr.numPr)

                    text = para.text.strip()

                    # Check for footnotes in the paragraph
                    footnotes = []
                    if hasattr(para._element, "xpath"):
                        for footnote_ref in para._element.xpath(
                            ".//w:footnoteReference"
                        ):
                            footnote_id = footnote_ref.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"
                            )
                            if hasattr(doc, "_footnotes") and doc._footnotes:
                                for footnote in doc._footnotes:
                                    if footnote.get("id") == footnote_id:
                                        footnote_text = self._extract_footnote_text(
                                            footnote
                                        )
                                        footnotes.append(
                                            f"[^{footnote_id}]: {footnote_text}"
                                        )

                    # Check for comments in the paragraph
                    comments = []
                    if hasattr(para._element, "xpath"):
                        for comment_ref in para._element.xpath(".//w:commentReference"):
                            comment_id = comment_ref.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"
                            )
                            if hasattr(doc, "_comments") and doc._comments:
                                for comment in doc._comments:
                                    if comment.get("id") == comment_id:
                                        comment_text = self._extract_comment_text(
                                            comment
                                        )
                                        comments.append(comment_text)

                    if (
                        not text and not footnotes and not comments
                    ):  # Skip empty paragraphs
                        continue

                    # Handle different types of paragraphs
                    if (
                        heading_level > 0
                    ):  # Changed: Remove is_real_heading check to capture more headings
                        # Save accumulated content before creating new heading
                        if current_content:
                            stack[-1].content = "\n".join(current_content)
                            current_content = []

                        new_node = DocxNode(text, heading_level, node_type="heading")
                        new_node.properties.update(
                            {
                                "style": style_name,
                                "is_real_heading": str(is_real_heading),
                            }
                        )

                        # Find appropriate parent by checking levels
                        while len(stack) > 1:  # Keep at least root
                            if stack[-1].level < heading_level or (
                                stack[-1].level == heading_level
                                and stack[-1].properties.get("is_real_heading")
                                == "True"
                                and is_real_heading is False
                            ):
                                break
                            stack.pop()

                        stack[-1].add_child(new_node)
                        stack.append(new_node)
                        in_list = False

                    elif is_list:
                        list_text = self._extract_list_text(para)
                        if not in_list:
                            current_content.append("")  # Add blank line before list
                        current_content.append(list_text)
                        # Add footnotes and comments if any
                        current_content.extend(footnotes)
                        current_content.extend(comments)
                        in_list = True

                    else:
                        # Regular paragraph
                        para_text = self._extract_paragraph_text(para)
                        if in_list:
                            current_content.append("")  # Add blank line after list
                        current_content.append(para_text)
                        # Add footnotes and comments if any
                        current_content.extend(footnotes)
                        current_content.extend(comments)
                        in_list = False

                elif element.tag.endswith("tbl"):  # Table
                    if current_content:
                        current_content.append("")  # Add blank line before table

                    from docx.table import Table

                    table = Table(element, doc)
                    table_text = self._extract_table_text(table)
                    current_content.append(table_text)
                    current_content.append("")  # Add blank line after table
                    in_list = False

                elif element.tag.endswith("sectPr"):  # Section break
                    if current_content:
                        stack[-1].content = "\n".join(current_content)
                        current_content = []
                    current_content.append(
                        "\n---\n"
                    )  # Add horizontal rule for section breaks
            except Exception as e:
                print(f"Error processing element: {e}")
                continue
        # Handle remaining content
        if current_content:
            stack[-1].content = "\n".join(current_content)

        # Clean up empty nodes
        def clean_node(node: DocxNode) -> bool:
            """Remove empty nodes and return whether this node should be kept"""
            # First clean children
            node.children = [child for child in node.children if clean_node(child)]
            # Keep node if it has content, non-empty children, or a meaningful title
            return bool(
                node.content.strip()
                or node.children
                or (node.title and node.title != "root")
            )

        root.children = [child for child in root.children if clean_node(child)]

        print(f"Document tree built with {len(root.children)} top-level sections")
        return root

    def _convert_to_outputs(
        self,
        node: DocxNode,
        id: str,
        parent_id: str = None,
        parent_titles: List[str] = None,
        parent_contents: List[str] = None,
        node_map: Dict[str, Node] = None,
    ) -> Tuple[List[Output], Dict[DocxNode, Output], SubGraph]:
        outputs = []
        node_chunk_map = {}
        nodes = []  # For SubGraph
        edges = []  # For SubGraph
        if node_map is None:
            node_map = {}  # Track title nodes by id
        chunk_nodes = {}  # Track chunk nodes by id

        if parent_titles is None:
            parent_titles = []
        if parent_contents is None:
            parent_contents = []

        # Use truncated title for path building
        truncated_title = (
            node.title[:50] + "..." if len(node.title) > 50 else node.title
        )
        current_titles = parent_titles + (
            [truncated_title] if node.title != "root" else []
        )
        print(f"\nProcessing node: {truncated_title}")
        print(f"Current titles path: {' / '.join(current_titles)}")

        def add_bidirectional_edge(
            from_node: Node, to_node: Node, label: str, properties: Dict = None
        ):
            """Helper function to add bidirectional edges between nodes"""
            if properties is None:
                properties = {}
            edges.append(Edge("", from_node, to_node, label, properties))
            edges.append(Edge("", to_node, from_node, label, properties))

        # Create title node for current node if it has a title
        if node.title != "root":
            # Use node's get_node_id method for stable ID generation
            title_node_id = node.get_node_id()
            if title_node_id not in node_map:
                title_node = Node(
                    _id=title_node_id,
                    name=truncated_title,  # Use truncated title for display
                    label="Title",
                    properties={
                        "level": str(node.level),
                        "type": node.node_type,
                        "full_title": node.title,  # Store full title as property
                        **node.properties,  # Include other properties
                    },
                )
                nodes.append(title_node)
                node_map[title_node_id] = title_node
                print(f"Created title node: {truncated_title}")

            # Create chunk for current node if it has content or is a leaf node
            if node.content or not node.children:
                # Use truncated titles for path
                full_title = " / ".join(filter(None, current_titles))
                print(f"Creating chunk for: {full_title}")

                # Store parent content separately
                parent_content = (
                    "\n".join(filter(None, parent_contents))
                    if parent_contents
                    else None
                )

                # Use node's content if available, otherwise use title as content for leaf nodes
                chunk_content = node.content if node.content else node.title
                current_output = Chunk(
                    id=f"{generate_hash_id(node.get_node_id())}",  # Use stable node ID for chunk ID
                    parent_id=parent_id,
                    name=full_title,
                    content=chunk_content,
                    parent_content=parent_content,
                    type=ChunkTypeEnum.Text,
                )

                # Apply length splitter if configured and content is too long
                if self.length_splitter:
                    split_chunks = self.length_splitter.slide_window_chunk(
                        current_output,
                        self.length_splitter.split_length,
                        self.length_splitter.window_length,
                    )
                    for idx, split_chunk in enumerate(split_chunks):
                        split_chunk.parent_id = current_output.id
                        # Add index to split chunk name for clarity
                        split_chunk.name = f"{full_title} (Part {idx + 1})"
                        outputs.append(split_chunk)
                        # Create chunk node for each split chunk
                        chunk_node = Node(
                            _id=split_chunk.id,
                            name=split_chunk.name,
                            label="Chunk",
                            properties={
                                "content": split_chunk.content[:100] + "..."
                            },  # Truncate content in properties
                        )
                        nodes.append(chunk_node)
                        chunk_nodes[split_chunk.id] = chunk_node
                        # Add bidirectional edges between title and split chunk
                        add_bidirectional_edge(
                            node_map[title_node_id], chunk_node, "hasContent"
                        )
                    # Map the original node to all split chunks
                    node_chunk_map[node] = split_chunks
                else:
                    outputs.append(current_output)
                    node_chunk_map[node] = current_output
                    # Create single chunk node
                    chunk_node = Node(
                        _id=current_output.id,
                        name=full_title,
                        label="Chunk",
                        properties={
                            "content": (
                                chunk_content[:100] + "..."
                                if len(chunk_content) > 100
                                else chunk_content
                            )
                        },
                    )
                    nodes.append(chunk_node)
                    chunk_nodes[current_output.id] = chunk_node
                    # Add bidirectional edges between title and chunk
                    add_bidirectional_edge(
                        node_map[title_node_id], chunk_node, "hasContent"
                    )

        # Process children
        if node.children:
            print(f"Processing {len(node.children)} children")
            for child in node.children:
                (
                    child_outputs,
                    child_node_chunk_map,
                    child_subgraph,
                ) = self._convert_to_outputs(
                    child,
                    id,
                    parent_id=parent_id,
                    parent_titles=current_titles,
                    parent_contents=parent_contents
                    + ([node.content] if node.content else []),
                    node_map=node_map,
                )
                outputs.extend(child_outputs)
                node_chunk_map.update(child_node_chunk_map)
                nodes.extend(child_subgraph.nodes)
                edges.extend(child_subgraph.edges)

                # Add title hierarchy relationship if both parent and child have titles
                if node.title != "root" and child.title != "root":
                    parent_node = node_map[node.get_node_id()]
                    child_node = node_map[child.get_node_id()]
                    add_bidirectional_edge(
                        parent_node, child_node, "hasChild", {"level": str(child.level)}
                    )

        return outputs, node_chunk_map, SubGraph(nodes=nodes, edges=edges)

    def _invoke(self, input: Input, **kwargs) -> Tuple[List[Output], SubGraph]:
        """
        Processes the input Docx file, extracts its text content, and generates Chunk objects.

        Args:
            input (Input): The file path of the Docx file to be processed.
            **kwargs: Additional keyword arguments, not used in the current implementation.

        Returns:
            Tuple[List[Output], SubGraph]: A tuple containing:
                - A list of Chunk objects based on the document structure
                - A SubGraph representing the document's title hierarchy

        Raises:
            ValueError: If the input is empty.
            IOError: If the file cannot be read or the text extraction fails.
        """
        if not input:
            raise ValueError("Input cannot be empty")

        try:
            doc = Document(input)
            document_root = self._build_document_tree(doc)
        except OSError as e:
            raise IOError(f"Failed to read file: {input}") from e

        basename, _ = os.path.splitext(os.path.basename(input))
        # Create a new root node with the file name and transfer document_root's children
        file_root = DocxNode(basename, 0)
        file_root.children = document_root.children  # Transfer children directly
        file_root.content = document_root.content  # Transfer content if any

        chunks, _, subgraph = self._convert_to_outputs(file_root, input)
        print("\nExtracted chunks:")
        for chunk in chunks:
            print(f"\nChunk: {chunk.name}")
            print(f"Content: {chunk.content}")
            if chunk.parent_content:
                print(f"Parent Content: {chunk.parent_content}")
        return chunks


if __name__ == "__main__":
    reader = ReaderABC.from_config({"type": "docx_reader"})
    chunks = reader.invoke(
        "/Users/zhangxinhong.zxh/Downloads/default.docx", write_ckpt=False
    )
    print("Extracted chunks:")
    for chunk in chunks:
        print(f"\nChunk: {chunk.name}")
        print(f"Content: {chunk.content}")
        if chunk.parent_content:
            print(f"Parent Content: {chunk.parent_content}")
